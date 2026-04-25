"""
PDF and document text extraction utilities.
"""

import io
from PyPDF2 import PdfReader


def extract_text_from_pdf(file) -> str:
    """
    Extract text from a PDF file.

    Args:
        file: A file-like object (e.g., Streamlit UploadedFile) or file path string.

    Returns:
        Extracted text as a string.
    """
    try:
        if isinstance(file, str):
            reader = PdfReader(file)
        else:
            # Handle Streamlit UploadedFile or BytesIO
            reader = PdfReader(io.BytesIO(file.read()))

        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"

        return text.strip()

    except Exception as e:
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")


def extract_text_from_docx(file) -> str:
    """
    Extract text from a DOCX file.

    Args:
        file: A file-like object or file path string.

    Returns:
        Extracted text as a string.
    """
    try:
        from docx import Document

        if isinstance(file, str):
            doc = Document(file)
        else:
            doc = Document(io.BytesIO(file.read()))

        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

        return text.strip()

    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")


def extract_text(file, filename: str = "") -> str:
    """
    Extract text from uploaded file based on file type.

    Supports: PDF, DOCX, TXT
    """
    if filename.lower().endswith(".pdf") or (hasattr(file, "name") and file.name.lower().endswith(".pdf")):
        return extract_text_from_pdf(file)
    elif filename.lower().endswith(".docx") or (hasattr(file, "name") and file.name.lower().endswith(".docx")):
        return extract_text_from_docx(file)
    elif filename.lower().endswith(".txt") or (hasattr(file, "name") and file.name.lower().endswith(".txt")):
        if isinstance(file, str):
            with open(file, "r") as f:
                return f.read().strip()
        else:
            return file.read().decode("utf-8").strip()
    else:
        # Try to read as plain text
        if hasattr(file, "read"):
            content = file.read()
            if isinstance(content, bytes):
                return content.decode("utf-8").strip()
            return content.strip()
        return str(file).strip()